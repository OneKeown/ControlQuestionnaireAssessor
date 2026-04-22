from io import BytesIO
from docx import Document


class ReportService:
    def build_docx_report(self, overall_summary: dict, control_results: list, certificate_results: list | None = None) -> BytesIO:
        doc = Document()

        doc.add_heading("Security Controls Assessment Report", level=1)

        doc.add_heading("Overall Assessment", level=2)
        doc.add_paragraph(f"Overall Result: {overall_summary['overall_status']}")
        doc.add_paragraph(f"Passed: {overall_summary['passed']} / {overall_summary['total']}")
        doc.add_paragraph(f"Failed: {overall_summary['failed']} / {overall_summary['total']}")
        doc.add_paragraph(f"Needs Review: {overall_summary['needs_review']} / {overall_summary['total']}")
        doc.add_paragraph(f"Summary: {overall_summary['summary_reason']}")

        if certificate_results:
            doc.add_heading("Certificate Checks", level=2)
            for cert in certificate_results:
                doc.add_paragraph(
                    f"{cert.get('file_name', 'Certificate')} | "
                    f"Type: {cert.get('cert_type', 'Unknown')} | "
                    f"Expiry: {cert.get('expiry_date', 'Not found')} | "
                    f"Status: {cert.get('status', 'Needs review')}"
                )

        doc.add_heading("Control Results", level=2)

        table = doc.add_table(rows=1, cols=6)
        hdr = table.rows[0].cells
        hdr[0].text = "Control ID"
        hdr[1].text = "Category"
        hdr[2].text = "Requirement"
        hdr[3].text = "Status"
        hdr[4].text = "Confidence"
        hdr[5].text = "Reason"

        for result in control_results:
            row = table.add_row().cells
            row[0].text = result.control_id
            row[1].text = result.category
            row[2].text = result.requirement
            row[3].text = result.status
            row[4].text = f"{result.confidence:.2f}"
            row[5].text = result.reason

        doc.add_heading("Evidence", level=2)
        for result in control_results:
            doc.add_paragraph(f"{result.control_id} — {result.category}", style="List Bullet")
            doc.add_paragraph(f"Status: {result.status}")
            if result.source_file:
                doc.add_paragraph(f"Source: {result.source_file} | Page: {result.source_page}")
            if result.source_excerpt:
                doc.add_paragraph(f"Excerpt: {result.source_excerpt}")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer