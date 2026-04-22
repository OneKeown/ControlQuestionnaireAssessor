import fitz
from docx import Document
from models.document_models import ExtractedDocument, ExtractedPage
from openpyxl import load_workbook


class ExtractorService:
    def extract_pdf(self, uploaded_file) -> ExtractedDocument:
        uploaded_file.seek(0)
        pdf_bytes = uploaded_file.read()
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")

        pages = []
        all_text = []

        for i, page in enumerate(doc):
            text = page.get_text("text")
            pages.append(ExtractedPage(page_number=i + 1, text=text))
            all_text.append(text)

        return ExtractedDocument(
            file_name=uploaded_file.name,
            file_type="pdf",
            pages=pages,
            full_text="\n".join(all_text),
        )

    def extract_txt(self, uploaded_file) -> ExtractedDocument:
        uploaded_file.seek(0)
        text = uploaded_file.read().decode("utf-8", errors="ignore")
        return ExtractedDocument(
            file_name=uploaded_file.name,
            file_type="txt",
            pages=[ExtractedPage(page_number=1, text=text)],
            full_text=text,
        )
    
    def extract_docx(self, uploaded_file) -> ExtractedDocument:
        uploaded_file.seek(0)
        doc = Document(uploaded_file)

        text_parts = []

        # paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())

        # tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        full_text = "\n".join(text_parts)

        return ExtractedDocument(
            file_name=uploaded_file.name,
            file_type="docx",
            pages=[
                ExtractedPage(page_number=1, text=full_text)
            ],
            full_text=full_text,
        )

    def extract_xlsx(self, uploaded_file) -> ExtractedDocument:
        uploaded_file.seek(0)

        workbook = load_workbook(uploaded_file, data_only=True)
        sheet_text_parts = []

        page_number = 1
        pages = []

        for sheet in workbook.worksheets:
            lines = [f"Sheet: {sheet.title}"]

            for row in sheet.iter_rows(values_only=True):
                values = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
                if values:
                    lines.append(" | ".join(values))

            sheet_text = "\n".join(lines).strip()

            if sheet_text:
                pages.append(
                    ExtractedPage(
                        page_number=page_number,
                        text=sheet_text
                    )
                )
                sheet_text_parts.append(sheet_text)
                page_number += 1

        full_text = "\n\n".join(sheet_text_parts)

        return ExtractedDocument(
            file_name=uploaded_file.name,
            file_type="xlsx",
            pages=pages if pages else [ExtractedPage(page_number=1, text="")],
            full_text=full_text,
        )

    def extract(self, uploaded_file) -> ExtractedDocument:
        name = uploaded_file.name.lower()

        if name.endswith(".pdf"):
            return self.extract_pdf(uploaded_file)

        if name.endswith(".txt"):
            return self.extract_txt(uploaded_file)

        if name.endswith(".docx"):  
            return self.extract_docx(uploaded_file)
        
        if name.endswith(".xlsx") or name.endswith(".xlsm"):
            return self.extract_xlsx(uploaded_file)


        raise ValueError(f"Unsupported file type: {uploaded_file.name}")